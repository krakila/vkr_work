"""Desktop client for the client-server stress analysis IS.

Tkinter is used for cross-platform GUI. The client waits for an active game
window, records screen video, captures input telemetry when pynput is available,
saves profiles to the server, and provides separate views for players and
researchers.
"""
from __future__ import annotations

import json
import math
import os
import queue
import re
import threading
import time
import wave
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

import cv2
import numpy as np
import requests
from PIL import Image, ImageGrab, ImageTk
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

try:
    import mss
except Exception:  # pragma: no cover
    mss = None

try:
    import pygetwindow as gw
except Exception:  # pragma: no cover
    gw = None

try:
    import serial
    from serial.tools import list_ports
except Exception:  # pragma: no cover
    serial = None
    list_ports = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    from pynput import keyboard as pynput_keyboard, mouse as pynput_mouse
except Exception:  # pragma: no cover
    pynput_keyboard = None
    pynput_mouse = None

from shared import (
    APP_NAME,
    ROLE_PLAYER,
    ROLE_RESEARCHER,
    active_window_title,
    classify_disease_group,
    combine_stress_scores,
    load_rule_calibration,
    movement_reversal,
    disease_codes_for_profile,
    ensure_dirs,
    heart_stress_score,
    looks_like_game_window,
    normalize_disease_group,
    payload_from_group_choice,
    profile_from_group_choice,
    role_label,
    sample_summary_text,
    safe_json_loads,
    stress_class,
    telemetry_stress_score,
    telemetry_summary,
    validate_password,
    validate_username,
    wait_for_game_window,
    slugify,
)

ensure_dirs()


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None:
            return default
        return float(value)
    except Exception:
        return default


def _fmt(value: Any, digits: int = 2, none: str = "—") -> str:
    if value is None:
        return none
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return none


APP_DIR = Path(__file__).resolve().parent
CACHE_DIR = APP_DIR / "data" / "cache"
CACHE_DIR.mkdir(parents=True, exist_ok=True)
DEFAULT_API_URL = os.environ.get("STRESS_IS_API", "http://127.0.0.1:8000")


class ApiClient:
    def __init__(self, base_url: str):
        self.base_url = base_url.rstrip("/")
        self.session = requests.Session()
        self.token = ""
        self.user: Dict[str, Any] = {}

    def set_auth(self, token: str, user: Dict[str, Any]):
        self.token = token or ""
        self.user = user or {}
        self.session.headers.update({"Authorization": f"Bearer {self.token}"} if self.token else {})

    def request(self, method: str, path: str, **kwargs):
        url = self.base_url + path
        resp = self.session.request(method, url, timeout=60, **kwargs)
        if resp.status_code >= 400:
            detail = None
            try:
                payload = resp.json()
                detail = payload.get("detail") or payload.get("message")
            except Exception:
                detail = None
            if not detail:
                detail = resp.text.strip() or f"HTTP {resp.status_code}"
            err = requests.HTTPError(detail)
            err.response = resp
            raise err
        return resp

    def get(self, path: str, **kwargs):
        return self.request("GET", path, **kwargs)

    def post(self, path: str, **kwargs):
        return self.request("POST", path, **kwargs)

    def put(self, path: str, **kwargs):
        return self.request("PUT", path, **kwargs)


class TelemetryCollector:
    """Global input telemetry with motion segmentation and bounded buffers."""

    def __init__(self):
        self.events: List[Dict[str, Any]] = []
        self._lock = threading.Lock()
        self._running = False
        self._mouse_listener = None
        self._keyboard_listener = None
        self._pressed_at: Dict[str, float] = {}
        self._last_mouse: Optional[Dict[str, Any]] = None
        self._last_mouse_emit_ts = 0.0
        self._motion_bucket: Optional[Dict[str, Any]] = None
        self._last_motion_flush_ts = 0.0
        self._last_segment_speed = 0.0
        self._last_segment_accel = 0.0
        self._last_click_ts = 0.0
        self._last_scroll_ts = 0.0
        self._last_pause_ts = 0.0
        self._max_age_sec = 35.0
        self._max_events = 3500

    def start(self):
        self.events.clear()
        self._pressed_at.clear()
        self._last_mouse = None
        self._last_mouse_emit_ts = 0.0
        self._motion_bucket = None
        self._last_motion_flush_ts = 0.0
        self._last_segment_speed = 0.0
        self._last_segment_accel = 0.0
        self._last_click_ts = 0.0
        self._last_scroll_ts = 0.0
        self._last_pause_ts = 0.0
        self._running = True
        if pynput_mouse and pynput_keyboard:
            self._start_global()
        return self

    def _prune_locked(self, now: Optional[float] = None):
        now = time.time() if now is None else now
        cutoff = now - self._max_age_sec
        if len(self.events) <= self._max_events and (not self.events or self.events[0].get("t", now) >= cutoff):
            return
        idx = 0
        for i, event in enumerate(self.events):
            if float(event.get("t", now)) >= cutoff:
                idx = i
                break
        else:
            idx = len(self.events)
        if idx > 0:
            self.events = self.events[idx:]
        if len(self.events) > self._max_events:
            self.events = self.events[-self._max_events:]

    def _append(self, event: Dict[str, Any]):
        with self._lock:
            self.events.append(event)
            self._prune_locked(float(event.get("t", time.time())))

    def _flush_motion_bucket(self, now: float, force: bool = False):
        if not self._motion_bucket:
            return
        bucket = self._motion_bucket
        duration = max(now - float(bucket["start_t"]), 1e-3)
        if not force and duration < 0.06 and float(bucket["path"]) < 18.0 and int(bucket["samples"]) < 4:
            return

        path = float(bucket["path"])
        net_dx = float(bucket["net_dx"])
        net_dy = float(bucket["net_dy"])
        net_dist = float((net_dx * net_dx + net_dy * net_dy) ** 0.5)
        straightness = net_dist / max(path, 1e-6)
        speed = path / max(duration, 1e-6)
        net_speed = net_dist / max(duration, 1e-6)
        accel = (speed - float(self._last_segment_speed)) / max(duration, 1e-6)
        jerk = (accel - float(self._last_segment_accel)) / max(duration, 1e-6)
        motion_event = {
            "t": now,
            "type": "mouse_move",
            "x": float(bucket["last_x"]),
            "y": float(bucket["last_y"]),
            "dx": net_dx,
            "dy": net_dy,
            "dist": path,
            "path": path,
            "net_dist": net_dist,
            "duration": duration,
            "speed": speed,
            "net_speed": net_speed,
            "accel": accel,
            "jerk": jerk,
            "turns": float(bucket["turns"]),
            "samples": int(bucket["samples"]),
            "straightness": straightness,
            "pause_after": max(0.0, now - float(bucket["last_t"])),
        }
        self.events.append(motion_event)
        self._prune_locked(now)
        self._motion_bucket = None
        self._last_segment_speed = speed
        self._last_segment_accel = accel
        self._last_motion_flush_ts = now

    def _start_global(self):
        def on_click(x, y, button, pressed):
            if not self._running or not pressed:
                return
            now = time.time()
            etype = "click"
            if self._last_click_ts and now - self._last_click_ts <= 0.25:
                etype = "rapid_click"
            self._last_click_ts = now
            self._append({"t": now, "type": etype, "x": x, "y": y, "button": str(button)})

        def on_scroll(x, y, dx, dy):
            if not self._running:
                return
            now = time.time()
            etype = "scroll"
            if self._last_scroll_ts and now - self._last_scroll_ts <= 0.25:
                etype = "scroll_burst"
            self._last_scroll_ts = now
            self._append({"t": now, "type": etype, "x": x, "y": y, "dx": dx, "dy": dy})

        def on_move(x, y):
            if not self._running:
                return
            now = time.time()
            if self._last_mouse is None:
                self._last_mouse = {"t": now, "x": x, "y": y}
                self._motion_bucket = {
                    "start_t": now,
                    "last_t": now,
                    "start_x": x,
                    "start_y": y,
                    "last_x": x,
                    "last_y": y,
                    "net_dx": 0.0,
                    "net_dy": 0.0,
                    "path": 0.0,
                    "samples": 0,
                    "turns": 0,
                    "prev_dx": 0.0,
                    "prev_dy": 0.0,
                }
                self._last_mouse_emit_ts = now
                return

            prev = self._last_mouse
            dt = max(now - float(prev["t"]), 1e-3)
            dx = float(x - float(prev["x"]))
            dy = float(y - float(prev["y"]))
            dist = float((dx * dx + dy * dy) ** 0.5)

            # Микродрожание и субпиксельные шаги не считаются отдельными событиями.
            if dt < 0.015 and dist < 1.15:
                self._last_mouse = {"t": now, "x": x, "y": y}
                return

            if self._motion_bucket is None:
                self._motion_bucket = {
                    "start_t": prev["t"],
                    "last_t": now,
                    "start_x": prev["x"],
                    "start_y": prev["y"],
                    "last_x": x,
                    "last_y": y,
                    "net_dx": dx,
                    "net_dy": dy,
                    "path": dist,
                    "samples": 1,
                    "turns": 0,
                    "prev_dx": dx,
                    "prev_dy": dy,
                }
            else:
                bucket = self._motion_bucket
                bucket["net_dx"] = float(bucket["net_dx"]) + dx
                bucket["net_dy"] = float(bucket["net_dy"]) + dy
                bucket["path"] = float(bucket["path"]) + dist
                bucket["samples"] = int(bucket["samples"]) + 1
                if movement_reversal(float(bucket["prev_dx"]), float(bucket["prev_dy"]), dx, dy, min_mag=10.0, cosine_threshold=-0.45):
                    bucket["turns"] = int(bucket["turns"]) + 1
                bucket["prev_dx"] = dx
                bucket["prev_dy"] = dy
                bucket["last_t"] = now
                bucket["last_x"] = x
                bucket["last_y"] = y

            self._last_mouse = {"t": now, "x": x, "y": y}

            bucket = self._motion_bucket
            duration = now - float(bucket["start_t"])
            if duration >= 0.09 or float(bucket["path"]) >= 30.0:
                self._flush_motion_bucket(now)
            elif now - self._last_mouse_emit_ts >= 0.16:
                # Редкая принудительная публикация помогает не терять длинные плавные движения.
                self._last_mouse_emit_ts = now
                self._flush_motion_bucket(now)

        def on_press(key):
            if not self._running:
                return
            name = getattr(key, "char", None) or str(key)
            lowered = name.lower()
            etype = "key_down"
            if "backspace" in lowered:
                etype = "backspace"
            elif "esc" in lowered:
                etype = "escape"
            self._pressed_at[name] = time.time()
            self._append({"t": time.time(), "type": etype, "key": name})

        def on_release(key):
            if not self._running:
                return
            name = getattr(key, "char", None) or str(key)
            hold = max(0.0, time.time() - self._pressed_at.pop(name, time.time()))
            etype = "key_up"
            if hold < 0.12:
                self._append({"t": time.time(), "type": "correction", "key": name, "hold": hold})
            self._append({"t": time.time(), "type": etype, "key": name, "hold": hold})

        try:
            self._mouse_listener = pynput_mouse.Listener(on_click=on_click, on_scroll=on_scroll, on_move=on_move)
            self._keyboard_listener = pynput_keyboard.Listener(on_press=on_press, on_release=on_release)
            self._mouse_listener.start()
            self._keyboard_listener.start()
        except Exception:
            self._mouse_listener = None
            self._keyboard_listener = None

    def stop(self):
        self._running = False
        now = time.time()
        try:
            self._flush_motion_bucket(now, force=True)
        except Exception:
            pass
        for listener in (self._mouse_listener, self._keyboard_listener):
            try:
                if listener:
                    listener.stop()
            except Exception:
                pass

    def snapshot(self) -> List[Dict[str, Any]]:
        now = time.time()
        with self._lock:
            try:
                self._flush_motion_bucket(now, force=True)
            except Exception:
                pass
            self._prune_locked(now)
            return list(self.events)



class HeartSensorReader:
    """Чтение данных Pulse Sensor из USB/COM-датчика в отдельном потоке.

    Датчик может отдавать либо строку вида "IBI: 842 BPM: 71", либо только BPM.
    При наличии IBI расчёты HRV строятся по межударным интервалам.
    """

    def __init__(self):
        self._lock = threading.Lock()
        self._running = False
        self._thread: Optional[threading.Thread] = None
        self.serial_port = None
        self.port = ""
        self.last_hr: Optional[float] = None
        self.last_ibi: Optional[float] = None
        self.last_bpm: Optional[float] = None
        self.last_line = ""
        self.last_error = ""
        self._ibi_samples: List[float] = []
        self._bpm_samples: List[float] = []

    def available_ports(self) -> List[str]:
        if list_ports is None:
            return []
        return [p.device for p in list(list_ports.comports())]

    def connect(self, port: str, baudrate: int = 115200) -> bool:
        if serial is None:
            self.last_error = "pyserial не установлен"
            return False
        if not port:
            self.last_error = "Не выбран COM/USB-порт"
            return False
        try:
            self.serial_port = serial.Serial(port, baudrate=baudrate, timeout=1)
            time.sleep(1.5)
            self.port = port
            self.last_error = ""
            self._running = True
            self._thread = threading.Thread(target=self._loop, daemon=True)
            self._thread.start()
            return True
        except Exception as e:
            self.last_error = str(e)
            self.serial_port = None
            return False

    def disconnect(self):
        self._running = False
        try:
            if self.serial_port is not None:
                self.serial_port.close()
        except Exception:
            pass
        self.serial_port = None
        self.port = ""

    @staticmethod
    def _parse_pulse_line(line: str) -> Dict[str, Optional[float]]:
        """Parse Pulse Sensor serial output.

        Supported patterns:
          - "IBI: 842 BPM: 71"
          - "BPM: 71"
          - "842,71"
          - "IBI=842;BPM=71"
        """
        line = (line or "").strip()
        result = {"ibi_ms": None, "bpm": None}
        if not line:
            return result

        # Prefer labeled values when present.
        m = re.search(r'\bIBI\s*[:=]\s*([-+]?\d+(?:[\.,]\d+)?)', line, re.I)
        if m:
            try:
                result["ibi_ms"] = float(m.group(1).replace(",", "."))
            except Exception:
                pass
        m = re.search(r'\bBPM\s*[:=]\s*([-+]?\d+(?:[\.,]\d+)?)', line, re.I)
        if m:
            try:
                result["bpm"] = float(m.group(1).replace(",", "."))
            except Exception:
                pass

        # Fallback for comma/space separated pairs.
        if result["ibi_ms"] is None and result["bpm"] is None:
            nums = []
            for token in re.findall(r'[-+]?\d+(?:[\.,]\d+)?', line):
                try:
                    nums.append(float(token.replace(",", ".")))
                except Exception:
                    continue
            if len(nums) >= 2:
                # Pulse Sensor serial logs usually print IBI first and BPM second.
                first, second = nums[0], nums[1]
                if first >= 250.0:
                    result["ibi_ms"] = first
                    result["bpm"] = second if 25.0 <= second <= 240.0 else None
                else:
                    # If the first number is too small for IBI, treat it as BPM.
                    result["bpm"] = first if 25.0 <= first <= 240.0 else None
                    result["ibi_ms"] = second if second >= 250.0 else None
            elif len(nums) == 1:
                value = nums[0]
                if 25.0 <= value <= 240.0:
                    result["bpm"] = value
                elif value >= 250.0:
                    result["ibi_ms"] = value

        # Reconstruct missing value when possible.
        if result["ibi_ms"] is None and result["bpm"] is not None and result["bpm"] > 0:
            result["ibi_ms"] = 60000.0 / result["bpm"]
        if result["bpm"] is None and result["ibi_ms"] is not None and result["ibi_ms"] > 0:
            result["bpm"] = 60000.0 / result["ibi_ms"]
        return result

    def _loop(self):
        while self._running and self.serial_port is not None:
            try:
                raw = self.serial_port.readline().decode('utf-8', errors='ignore').strip()
                if not raw:
                    continue
                parsed = self._parse_pulse_line(raw)
                with self._lock:
                    self.last_line = raw
                    if parsed["ibi_ms"] is not None:
                        self.last_ibi = float(parsed["ibi_ms"])
                        self._ibi_samples.append(self.last_ibi)
                        self._ibi_samples = self._ibi_samples[-256:]
                    if parsed["bpm"] is not None:
                        self.last_bpm = float(parsed["bpm"])
                        self._bpm_samples.append(self.last_bpm)
                        self._bpm_samples = self._bpm_samples[-256:]
                    if self.last_ibi is not None:
                        self.last_hr = 60000.0 / max(self.last_ibi, 1e-6)
                    elif self.last_bpm is not None:
                        self.last_hr = self.last_bpm
            except Exception as e:
                self.last_error = str(e)
                time.sleep(0.2)

    def sample(self) -> Optional[float]:
        with self._lock:
            return self.last_hr

    def sample_packet(self) -> Dict[str, Optional[float]]:
        with self._lock:
            return {
                "hr": self.last_hr,
                "ibi_ms": self.last_ibi,
                "bpm": self.last_bpm,
                "line": self.last_line,
            }

    def ibi_series(self) -> List[float]:
        with self._lock:
            return list(self._ibi_samples)

    def bpm_series(self) -> List[float]:
        with self._lock:
            return list(self._bpm_samples)

    def status_text(self) -> str:
        if self.serial_port is None:
            return "Не подключён"
        if self.last_ibi is not None:
            return f"Подключён: {self.port} | IBI: {self.last_ibi:.0f} мс | BPM: {self.last_bpm:.0f}" if self.last_bpm is not None else f"Подключён: {self.port} | IBI: {self.last_ibi:.0f} мс"
        if self.last_bpm is not None:
            return f"Подключён: {self.port} | BPM: {self.last_bpm:.0f}"
        return f"Подключён: {self.port}"


class ScreenRecorder:
    def __init__(self, fps: int = 10, capture_active_window: bool = True, blank_mode: bool = False):
        self.fps = max(1, int(fps))
        self.capture_active_window = capture_active_window
        self.blank_mode = blank_mode
        self.video_path: Optional[Path] = None
        self.audio_path: Optional[Path] = None
        self._running = False
        self._thread = None
        self._size_ready = threading.Event()
        self.width = 0
        self.height = 0
        self.start_ts = 0.0
        self.end_ts = 0.0
        self._writer = None
        self._last_frame_ts = 0.0
        self._cached_region = None
        self._cached_region_ts = 0.0
        self._max_width = 1280
        self._max_height = 720

    def _window_region(self):
        if not self.capture_active_window or self.blank_mode or gw is None:
            return None
        now = time.time()
        if self._cached_region is not None and (now - self._cached_region_ts) < 0.5:
            return self._cached_region
        try:
            win = gw.getActiveWindow()
            if not win:
                return None
            left = int(getattr(win, "left", 0) or 0)
            top = int(getattr(win, "top", 0) or 0)
            width = int(getattr(win, "width", 0) or 0)
            height = int(getattr(win, "height", 0) or 0)
            if width <= 0 or height <= 0:
                return None
            region = {"left": left, "top": top, "width": width, "height": height}
            self._cached_region = region
            self._cached_region_ts = now
            return region
        except Exception:
            return None

    def _grab_frame(self):
        if self.blank_mode:
            if self.width <= 0 or self.height <= 0:
                self.width, self.height = 1280, 720
            return np.full((self.height, self.width, 3), 255, dtype=np.uint8)
        region = self._window_region()
        if mss is not None and region is not None:
            try:
                with mss.mss() as sct:
                    return self._grab_frame_with_sct(sct)
            except Exception:
                pass
        try:
            img = ImageGrab.grab(all_screens=False)
            return cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
        except Exception:
            return None

    def _grab_frame_with_sct(self, sct):
        region = self._window_region()
        if region is None:
            return None
        try:
            shot = sct.grab(region)
            frame = np.array(shot)
            if frame.shape[-1] == 4:
                frame = frame[:, :, :3]
                return cv2.cvtColor(frame, cv2.COLOR_BGRA2BGR)
            return frame
        except Exception:
            return None

    def _resize_if_needed(self, frame):
        if frame is None:
            return None
        h, w = frame.shape[:2]
        if w <= self._max_width and h <= self._max_height:
            return frame
        scale = min(self._max_width / float(w), self._max_height / float(h))
        new_w = max(1, int(w * scale))
        new_h = max(1, int(h * scale))
        return cv2.resize(frame, (new_w, new_h), interpolation=cv2.INTER_AREA)

    def start(self, out_dir: Path):
        out_dir.mkdir(parents=True, exist_ok=True)
        self.video_path = out_dir / "session.mp4"
        self.audio_path = out_dir / "session.wav"
        self.start_ts = time.time()
        self._running = True
        self._thread = threading.Thread(target=self._run, daemon=True)
        self._thread.start()
        return self

    def _run(self):
        fourcc = cv2.VideoWriter_fourcc(*"mp4v")
        first = True
        frame_interval = 1.0 / float(self.fps)
        next_tick = time.perf_counter()

        # Не крутим плотный busy-loop: захват ограничен FPS и уменьшенным размером,
        # чтобы запись видео не мешала вводу мыши в игре.
        if mss is not None:
            try:
                with mss.mss() as sct:
                    while self._running:
                        frame = self._grab_frame_with_sct(sct)
                        frame = self._resize_if_needed(frame)
                        if frame is not None:
                            h, w = frame.shape[:2]
                            if first:
                                self.width, self.height = w, h
                                self._writer = cv2.VideoWriter(str(self.video_path), fourcc, float(self.fps), (w, h))
                                self._size_ready.set()
                                first = False
                            if self._writer:
                                try:
                                    self._writer.write(frame)
                                except Exception:
                                    pass
                        next_tick += frame_interval
                        sleep_for = max(0.0, next_tick - time.perf_counter())
                        if sleep_for > 0:
                            time.sleep(sleep_for)
            except Exception:
                pass
        else:
            while self._running:
                frame = self._grab_frame()
                frame = self._resize_if_needed(frame)
                if frame is not None:
                    h, w = frame.shape[:2]
                    if first:
                        self.width, self.height = w, h
                        self._writer = cv2.VideoWriter(str(self.video_path), fourcc, float(self.fps), (w, h))
                        self._size_ready.set()
                        first = False
                    if self._writer:
                        try:
                            self._writer.write(frame)
                        except Exception:
                            pass
                next_tick += frame_interval
                sleep_for = max(0.0, next_tick - time.perf_counter())
                if sleep_for > 0:
                    time.sleep(sleep_for)
        try:
            if self._writer:
                self._writer.release()
        except Exception:
            pass
        self.end_ts = time.time()

    def wait_ready(self, timeout: float = 5.0) -> bool:
        return self._size_ready.wait(timeout)

    def stop(self):
        self._running = False
        if self._thread:
            self._thread.join(timeout=3)
        return self.video_path, self.audio_path


def create_blank_video(out_path: Path, duration: float, fps: int = 15, size: tuple[int, int] = (1280, 720)) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    width, height = size
    frame_count = max(1, int(max(duration, 1.0) * fps))
    fourcc = cv2.VideoWriter_fourcc(*"mp4v")
    writer = cv2.VideoWriter(str(out_path), fourcc, float(fps), (width, height))
    frame = np.full((height, width, 3), 255, dtype=np.uint8)
    for _ in range(frame_count):
        writer.write(frame)
    writer.release()
    return out_path


def synth_audio_from_timeline(timeline: List[Dict[str, Any]], out_path: Path, duration: float, sample_rate: int = 22050):
    samples = max(1, int(duration * sample_rate))
    data = np.zeros(samples, dtype=np.float32)
    for item in timeline:
        t = float(item.get("t", 0.0))
        overall = float(item.get("overall_score", 0.0))
        pos = int(t * sample_rate)
        if 0 <= pos < samples:
            length = min(int(sample_rate * 0.08), samples - pos)
            freq = 180 + int(260 * overall)
            tvec = np.linspace(0.0, 0.08, length, endpoint=False)
            data[pos:pos + length] += (0.2 + 0.7 * overall) * np.sin(2 * np.pi * freq * tvec)
    data = np.clip(data, -1.0, 1.0)
    with wave.open(str(out_path), "wb") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sample_rate)
        wf.writeframes((data * 32767).astype(np.int16).tobytes())
    return out_path


class LoginFrame(ttk.Frame):
    def __init__(self, master, app):
        super().__init__(master, padding=18)
        self.app = app
        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=1)

        ttk.Label(self, text=APP_NAME, font=("Segoe UI", 19, "bold")).grid(row=0, column=0, columnspan=2, sticky="w")
        ttk.Label(self, text="Вход, регистрация и сохранение входа", font=("Segoe UI", 11)).grid(row=1, column=0, columnspan=2, sticky="w", pady=(0, 12))

        self.username = tk.StringVar()
        self.password = tk.StringVar()
        self.role = tk.StringVar(value=ROLE_PLAYER)
        self.remember = tk.BooleanVar(value=True)

        form = ttk.LabelFrame(self, text="Учётная запись", padding=12)
        form.grid(row=2, column=0, sticky="nsew", padx=(0, 10))
        form.columnconfigure(1, weight=1)
        ttk.Label(form, text="Логин").grid(row=0, column=0, sticky="w", pady=4)
        ttk.Entry(form, textvariable=self.username).grid(row=0, column=1, sticky="ew", pady=4)
        ttk.Label(form, text="Пароль").grid(row=1, column=0, sticky="w", pady=4)
        ttk.Entry(form, textvariable=self.password, show="*").grid(row=1, column=1, sticky="ew", pady=4)
        ttk.Label(form, text="Роль").grid(row=2, column=0, sticky="w", pady=4)
        ttk.Combobox(form, textvariable=self.role, values=[ROLE_PLAYER, ROLE_RESEARCHER], state="readonly").grid(row=2, column=1, sticky="ew", pady=4)
        ttk.Checkbutton(form, text="Сохранить вход", variable=self.remember).grid(row=3, column=0, columnspan=2, sticky="w", pady=(6, 10))
        buttons = ttk.Frame(form)
        buttons.grid(row=4, column=0, columnspan=2, sticky="ew")
        ttk.Button(buttons, text="Войти", command=self.login).pack(side="left", expand=True, fill="x", padx=(0, 6))
        ttk.Button(buttons, text="Зарегистрироваться", command=self.register).pack(side="left", expand=True, fill="x")

        hint = ttk.LabelFrame(self, text="Подсказка", padding=12)
        hint.grid(row=2, column=1, sticky="nsew")
        ttk.Label(
            hint,
            text=(
                "Игрок видит только свои данные.\n"
                "Исследователь видит всех игроков и сравнение по играм.\n\n"
                "Пустые логины и пароли запрещены."
            ),
            justify="left",
        ).pack(anchor="w")

    def login(self):
        try:
            validate_username(self.username.get())
            validate_password(self.password.get())
        except ValueError as e:
            messagebox.showerror("Ошибка", str(e))
            return
        self.app.do_login(self.username.get(), self.password.get())

    def register(self):
        try:
            validate_username(self.username.get())
            validate_password(self.password.get())
        except ValueError as e:
            messagebox.showerror("Ошибка", str(e))
            return
        self.app.do_register(self.username.get(), self.password.get(), self.role.get())


class LivePlayerTab(ttk.Frame):
    def __init__(self, master, app):
        super().__init__(master, padding=10)
        self.app = app
        self.session_id: Optional[int] = None
        self.running = False
        self.waiting = False
        self.collector = TelemetryCollector()
        self.heart_sensor = HeartSensorReader()
        self.hr_samples: List[float] = []
        self.ibi_samples: List[float] = []
        self.bpm_samples: List[float] = []
        self.recorder: Optional[ScreenRecorder] = None
        self.timeline: List[Dict[str, Any]] = []
        self.current_game_title = ""
        self.wait_deadline = 0.0
        self._poll_job = None
        self._ui_job = None

        left = ttk.Frame(self)
        left.pack(side="left", fill="y", padx=(0, 12))
        right = ttk.Frame(self)
        right.pack(side="left", fill="both", expand=True)

        profile = ttk.LabelFrame(left, text="Профиль игрока", padding=10)
        profile.pack(fill="x", pady=(0, 10))
        self.full_name = tk.StringVar(value=self.app.current_user.get("profile", {}).get("full_name", self.app.current_user.get("username", "")))
        self.age = tk.IntVar(value=int(self.app.current_user.get("profile", {}).get("age", 18) or 18))
        self.sex = tk.StringVar(value=(self.app.current_user.get("profile", {}).get("sex", "мужской") or "мужской"))
        self.disease_group = tk.StringVar(value=normalize_disease_group(self.app.current_user.get("profile", {}).get("disease_group", "healthy")))

        ttk.Label(profile, text="ФИО").grid(row=0, column=0, sticky="w")
        ttk.Entry(profile, textvariable=self.full_name, width=28).grid(row=0, column=1, sticky="ew", pady=2)
        ttk.Label(profile, text="Возраст").grid(row=1, column=0, sticky="w")
        ttk.Spinbox(profile, from_=1, to=120, textvariable=self.age, width=8).grid(row=1, column=1, sticky="w", pady=2)
        ttk.Label(profile, text="Пол").grid(row=2, column=0, sticky="w")
        ttk.Combobox(profile, textvariable=self.sex, values=["мужской", "женский"], state="readonly").grid(row=2, column=1, sticky="ew", pady=2)
        ttk.Label(profile, text="Состояние").grid(row=3, column=0, sticky="w")
        disease_row = ttk.Frame(profile)
        disease_row.grid(row=3, column=1, sticky="ew")
        for text, value in [("Здоровые", "healthy"), ("Астма / лёгкие", "asthma"), ("Сердечные", "heart")]:
            ttk.Radiobutton(disease_row, text=text, value=value, variable=self.disease_group).pack(anchor="w")
        ttk.Button(profile, text="Сохранить профиль", command=self.save_profile).grid(row=4, column=0, columnspan=2, sticky="ew", pady=(8, 0))
        profile.columnconfigure(1, weight=1)

        session_box = ttk.LabelFrame(left, text="Сессия", padding=10)
        session_box.pack(fill="x", pady=(0, 10))
        self.game_title = tk.StringVar(value="")
        self.notes = tk.StringVar(value="")
        self.analysis_mode = tk.StringVar(value="rule")
        self.record_video = tk.BooleanVar(value=True)
        ttk.Label(session_box, text="Название игры").grid(row=0, column=0, sticky="w")
        ttk.Entry(session_box, textvariable=self.game_title, width=28).grid(row=0, column=1, sticky="ew", pady=2)
        ttk.Label(session_box, text="Заметка").grid(row=1, column=0, sticky="w")
        ttk.Entry(session_box, textvariable=self.notes, width=28).grid(row=1, column=1, sticky="ew", pady=2)
        ttk.Label(session_box, text="Режим оценки").grid(row=2, column=0, sticky="w")
        mode_row = ttk.Frame(session_box)
        mode_row.grid(row=2, column=1, sticky="ew")
        ttk.Radiobutton(mode_row, text="Rule-based", value="rule", variable=self.analysis_mode).pack(anchor="w")
        ttk.Radiobutton(mode_row, text="ML", value="ml", variable=self.analysis_mode).pack(anchor="w")
        ttk.Checkbutton(session_box, text="Записывать видео", variable=self.record_video).grid(row=3, column=0, columnspan=2, sticky="w", pady=(8, 2))
        ttk.Label(session_box, text="При выключенной записи сохраняется белый видеофайл.", wraplength=230, foreground="#555").grid(row=4, column=0, columnspan=2, sticky="w", pady=(0, 6))
        self.start_button = ttk.Button(session_box, text="Начать с ожиданием игры", command=self.start_session)
        self.start_button.grid(row=5, column=0, columnspan=2, sticky="ew", pady=2)
        self.stop_button = ttk.Button(session_box, text="Стоп и отправить", command=self.stop_session)
        self.stop_button.grid(row=6, column=0, columnspan=2, sticky="ew", pady=2)
        ttk.Button(session_box, text="Обновить сессии", command=self.app.refresh_views).grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)
        session_box.columnconfigure(1, weight=1)

        sensor = ttk.LabelFrame(left, text="Датчик ЧСС (USB/COM)", padding=10)
        sensor.pack(fill="x", pady=(0, 10))
        self.hr_port = tk.StringVar(value="")
        self.hr_status = tk.StringVar(value=self.heart_sensor.status_text())
        ttk.Label(sensor, text="Порт").grid(row=0, column=0, sticky="w")
        self.port_box = ttk.Combobox(sensor, textvariable=self.hr_port, values=self.heart_sensor.available_ports(), state="readonly")
        self.port_box.grid(row=0, column=1, sticky="ew", pady=2)
        ttk.Button(sensor, text="Обновить", command=self.refresh_hr_ports).grid(row=1, column=0, sticky="ew", pady=2)
        ttk.Button(sensor, text="Подключить", command=self.connect_hr_sensor).grid(row=1, column=1, sticky="ew", pady=2)
        ttk.Label(sensor, textvariable=self.hr_status, wraplength=210).grid(row=2, column=0, columnspan=2, sticky="w", pady=(4, 0))
        sensor.columnconfigure(1, weight=1)

        self.status = tk.StringVar(value="Ожидание")
        ttk.Label(left, textvariable=self.status, font=("Segoe UI", 10, "bold")).pack(anchor="w", pady=(0, 8))

        metrics = ttk.LabelFrame(right, text="Текущие показатели", padding=10)
        metrics.pack(fill="x")
        self.live_vars = {k: tk.StringVar(value="—") for k in ["heart", "telemetry", "ml", "overall", "class", "source", "game"]}
        labels = [
            ("game", "Окно игры"), ("heart", "ЧСС / IBI / HRV"), ("telemetry", "Телеметрия"),
            ("overall", "Итог"), ("class", "Класс"), ("source", "Источник стресса"),
        ]
        for i, (k, caption) in enumerate(labels):
            ttk.Label(metrics, text=caption).grid(row=i // 2, column=(i % 2) * 2, sticky="w", padx=(0, 8), pady=3)
            ttk.Label(metrics, textvariable=self.live_vars[k]).grid(row=i // 2, column=(i % 2) * 2 + 1, sticky="w", pady=3)
        metrics.columnconfigure(1, weight=1)
        metrics.columnconfigure(3, weight=1)

        self.detail_box = ttk.LabelFrame(right, text="Последние события телеметрии", padding=8)
        self.detail_box.pack(fill="both", expand=False, pady=(10, 10))
        self.telemetry_list = tk.Listbox(self.detail_box, height=8)
        self.telemetry_list.pack(fill="both", expand=True)

        self.graph_frame = ttk.LabelFrame(right, text="Графики", padding=8)
        self.graph_frame.pack(fill="both", expand=True)
        self.fig = Figure(figsize=(7, 4), dpi=100)
        self.ax = self.fig.add_subplot(111)
        self.canvas = FigureCanvasTkAgg(self.fig, master=self.graph_frame)
        self.canvas.get_tk_widget().pack(fill="both", expand=True)

    def _current_profile_payload(self) -> Dict[str, Any]:
        return profile_from_group_choice(
            self.full_name.get().strip() or self.app.current_user.get("username", ""),
            self.age.get(),
            self.sex.get(),
            self.disease_group.get(),
            self.notes.get(),
        )

    def save_profile(self):
        try:
            profile = self._current_profile_payload()
            if not profile["full_name"]:
                raise ValueError("ФИО не может быть пустым")
            if int(profile["age"]) <= 0:
                raise ValueError("Возраст должен быть положительным")
            resp = self.app.api.post("/profile", json=profile)
            resp.raise_for_status()
            self.app.current_user["profile"] = resp.json()["profile"]
            self.status.set("Профиль сохранён")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить профиль: {e}")

    def refresh_hr_ports(self):
        if list_ports is None:
            self.hr_status.set("pyserial не установлен")
            return
        ports = [p.device for p in list(list_ports.comports())]
        self.port_box.configure(values=ports)
        if ports and self.hr_port.get() not in ports:
            self.hr_port.set(ports[0])
        self.hr_status.set("Порты обновлены" if ports else "COM/USB-порты не найдены")

    def connect_hr_sensor(self):
        port = self.hr_port.get().strip()
        if self.heart_sensor.serial_port is not None:
            self.heart_sensor.disconnect()
            self.hr_status.set("Датчик ЧСС отключён")
            return
        if self.heart_sensor.connect(port):
            self.hr_status.set(f"Датчик ЧСС подключён: {port}")
        else:
            self.hr_status.set(f"Не удалось подключить датчик: {self.heart_sensor.last_error}")

    def start_session(self):
        if self.running or self.waiting:
            return
        self.waiting = True
        self.wait_deadline = time.time() + 120.0
        self.status.set("Ожидание активного окна игры...")
        self._update_ui()
        self._poll_wait_for_window()

    def _poll_wait_for_window(self):
        if not self.waiting:
            return
        title = active_window_title()
        hint = self.game_title.get().strip()
        if title and looks_like_game_window(title):
            if not hint or hint.lower() in title.lower():
                self._begin_session(title)
                return
        if time.time() >= self.wait_deadline:
            self.waiting = False
            self._update_ui()
            self.status.set("Ожидание истекло")
            messagebox.showwarning("Ожидание игры", "Не удалось дождаться активного окна игры.")
            return
        if title:
            self.status.set(f"Ожидание игры: {title}")
        self.after(300, self._poll_wait_for_window)

    def _begin_session(self, active_title: str):
        self.waiting = False
        self.running = True
        self.current_game_title = active_title.strip() or self.game_title.get().strip() or "Игра"
        self.game_title.set(self.current_game_title)
        try:
            self.save_profile()
        except Exception:
            pass
        payload = {
            "game_title": self.current_game_title,
            "notes": self.notes.get().strip(),
            "full_name": self.full_name.get().strip() or self.app.current_user.get("username", ""),
            "age": int(self.age.get()),
            "sex": self.sex.get(),
            "disease_group": self.disease_group.get(),
            "disease_codes": disease_codes_for_profile(self._current_profile_payload(), 1),
        }
        try:
            session = self.app.api.post("/sessions", json=payload).json()
        except Exception as e:
            self.running = False
            messagebox.showerror("Ошибка", f"Не удалось создать сессию: {e}")
            return
        self.session_id = session["id"]
        self.timeline = []
        self.collector.start()
        self.recorder = ScreenRecorder(
            fps=15,
            capture_active_window=True,
            blank_mode=not self.record_video.get(),
        ).start(CACHE_DIR / f"session_{self.session_id}")
        if self.record_video.get():
            self.status.set(f"Запись идёт. Сессия №{self.session_id}")
        else:
            self.status.set(f"Сессия №{self.session_id} идёт без записи видео (белый экран)")
        self._poll_live()
        self._update_ui()

    def _update_ui(self):
        start_state = "disabled" if (self.running or self.waiting) else "normal"
        stop_state = "normal" if (self.running or self.waiting) else "disabled"
        for widget_name, state in (("start_button", start_state), ("stop_button", stop_state)):
            widget = getattr(self, widget_name, None)
            try:
                if widget is not None and widget.winfo_exists():
                    widget.configure(state=state)
            except Exception:
                pass

    def _select_recent_events(self, snapshot: List[Dict[str, Any]], window_sec: float = 10.0) -> List[Dict[str, Any]]:
        now = time.time()
        return [e for e in snapshot if now - float(e.get("t", now)) <= window_sec]

    def _mock_hr_series(self, t_rel: float) -> List[float]:
        base = 72.0 + 6.0 * math.sin(t_rel / 7.0)
        if self.timeline and len(self.timeline) % 15 == 0:
            base += 10.0
        if self.live_vars["class"].get() == "высокий":
            base += 8.0
        return [max(48.0, min(165.0, base + math.sin(t_rel / 2.5) * 2.0)) for _ in range(8)]

    def _analysis_ml_prob(self, features: Dict[str, float]) -> float:
        try:
            resp = self.app.api.post("/ml/predict", json=features)
            return float(resp.json().get("ml_prob", 0.5))
        except Exception:
            return combine_stress_scores(features.get("heart_score", 0.0), features.get("telemetry_score", 0.0), 0.5)

    def _poll_live(self):
        if not self.running:
            return
        snap = self.collector.snapshot()
        window = self._select_recent_events(snap, 10.0)
        start_ts = self.recorder.start_ts if self.recorder else (self.timeline[0]["t"] if self.timeline else time.time())
        t_rel = max(0.0, time.time() - start_ts)

        pulse_packet = self.heart_sensor.sample_packet()
        hr_value = pulse_packet.get("hr")
        ibi_value = pulse_packet.get("ibi_ms")
        bpm_value = pulse_packet.get("bpm")

        if hr_value is not None:
            self.hr_samples.append(float(hr_value))
            self.hr_samples = self.hr_samples[-256:]
        if ibi_value is not None:
            self.ibi_samples.append(float(ibi_value))
            self.ibi_samples = self.ibi_samples[-256:]
        if bpm_value is not None:
            self.bpm_samples.append(float(bpm_value))
            self.bpm_samples = self.bpm_samples[-256:]

        if len(self.ibi_samples) >= 2:
            ibi_arr = np.array(self.ibi_samples[-12:], dtype=float)
            hr_arr = 60000.0 / np.clip(ibi_arr, 1e-6, None)
            ibi_mean = float(np.mean(ibi_arr))
            ibi_std = float(np.std(ibi_arr, ddof=1)) if len(ibi_arr) > 1 else 0.0
            if len(ibi_arr) > 1:
                diff_ibi = np.diff(ibi_arr)
                sdnn_v = float(np.std(ibi_arr, ddof=1))
                rmssd_v = float(np.sqrt(np.mean(diff_ibi ** 2)))
                pnn50_v = float(np.mean(np.abs(diff_ibi) > 50.0))
            else:
                sdnn_v = rmssd_v = pnn50_v = 0.0

            rounded = np.round(ibi_arr / 50.0) * 50.0
            if len(rounded):
                unique, counts = np.unique(rounded.astype(int), return_counts=True)
                mode_ibi = float(unique[int(np.argmax(counts))]) if len(unique) else 0.0
                amo = float(100.0 * max(counts) / max(len(ibi_arr), 1)) if len(counts) else 0.0
                baevsky_v = float((amo * 100.0) / max(2.0 * max(mode_ibi, 1e-6) * max(np.ptp(ibi_arr), 1e-6), 1e-6)) if mode_ibi > 0 else 0.0
            else:
                baevsky_v = 0.0

            phys = {
                "hr_mean": float(np.mean(hr_arr)),
                "hr_std": float(np.std(hr_arr, ddof=1)) if len(hr_arr) > 1 else 0.0,
                "ibi_mean": ibi_mean,
                "ibi_std": ibi_std,
                "rr_mean": ibi_mean,
                "sdnn": sdnn_v,
                "rmssd": rmssd_v,
                "pnn50": pnn50_v,
                "baevsky": baevsky_v,
                "ibi_samples": len(ibi_arr),
                "bpm_mean": float(np.mean(self.bpm_samples[-12:])) if self.bpm_samples else 0.0,
            }
            heart = heart_stress_score(phys)
        elif len(self.hr_samples) >= 2:
            hr_arr = np.array(self.hr_samples[-12:], dtype=float)
            ibi_arr = 60000.0 / np.clip(hr_arr, 1e-6, None)
            ibi_mean = float(np.mean(ibi_arr))
            ibi_std = float(np.std(ibi_arr, ddof=1)) if len(ibi_arr) > 1 else 0.0
            if len(ibi_arr) > 1:
                diff_ibi = np.diff(ibi_arr)
                sdnn_v = float(np.std(ibi_arr, ddof=1))
                rmssd_v = float(np.sqrt(np.mean(diff_ibi ** 2)))
                pnn50_v = float(np.mean(np.abs(diff_ibi) > 50.0))
            else:
                sdnn_v = rmssd_v = pnn50_v = 0.0
            phys = {
                "hr_mean": float(np.mean(hr_arr)),
                "hr_std": float(np.std(hr_arr, ddof=1)) if len(hr_arr) > 1 else 0.0,
                "ibi_mean": ibi_mean,
                "ibi_std": ibi_std,
                "rr_mean": ibi_mean,
                "sdnn": sdnn_v,
                "rmssd": rmssd_v,
                "pnn50": pnn50_v,
                "baevsky": 0.0,
                "ibi_samples": len(ibi_arr),
                "bpm_mean": float(np.mean(self.bpm_samples[-12:])) if self.bpm_samples else 0.0,
            }
            heart = heart_stress_score(phys)
        else:
            phys = {"hr_mean": 0.0, "hr_std": 0.0, "ibi_mean": 0.0, "ibi_std": 0.0, "rr_mean": 0.0, "sdnn": 0.0, "rmssd": 0.0, "pnn50": 0.0, "baevsky": 0.0, "ibi_samples": 0, "bpm_mean": 0.0}
            heart = None

        tel = telemetry_summary(window, window_seconds=10.0)
        telemetry = telemetry_stress_score(tel)
        ml_prob = self._analysis_ml_prob({
            "hr_mean": phys["hr_mean"], "hr_std": phys["hr_std"], "sdnn": phys["sdnn"], "rmssd": phys["rmssd"],
            "pnn50": phys["pnn50"], "baevsky": phys["baevsky"],
            **tel,
        })
        if self.analysis_mode.get() == "ml":
            overall = ml_prob
        else:
            try:
                weights = self.app.rule_weights
                overall = combine_stress_scores(heart, telemetry, None, weights={"heart": weights.get("heart", 0.5), "telemetry": weights.get("telemetry", 0.5), "ml": 0.0})
            except Exception:
                overall = combine_stress_scores(heart, telemetry, None, weights={"heart": 0.5, "telemetry": 0.5, "ml": 0.0})
        row = {
            "t": round(t_rel, 2),
            "heart_score": heart,
            "telemetry_score": telemetry,
            "ml_score": ml_prob,
            "overall_score": overall,
            "class": stress_class(overall),
            "source": "ML" if self.analysis_mode.get() == "ml" else "Rule-based",
            **phys,
            **tel,
        }
        self.timeline.append(row)
        self._render_live(row)
        self._poll_job = self.after(350, self._poll_live)

    def _render_live(self, row: Dict[str, Any]):
        self.live_vars["game"].set(self.current_game_title)
        heart_text = "датчик не подключён" if row.get("heart_score") is None else f"{row['heart_score']:.2f} | ЧСС {row['hr_mean']:.0f} уд/мин | IBI {row.get('ibi_mean', 0.0):.0f} мс"
        self.live_vars["heart"].set(heart_text)
        self.live_vars["telemetry"].set(f"{row['telemetry_score']:.2f} | ошибки {row['error_rate']:.2f}, клики {row['click_rate']:.2f}")
        self.live_vars["ml"].set(f"{row['ml_score']:.2f} | ML" if self.analysis_mode.get() == "ml" else "—")
        self.live_vars["overall"].set(f"{row['overall_score']:.2f}")
        self.live_vars["class"].set(row["class"])
        self.live_vars["source"].set(f"{row['source']} | Ввод: {row['event_density']:.2f} событий/с")
        self._draw_graph()
        self.telemetry_list.delete(0, tk.END)
        for item in self.timeline[-30:]:
            heart_val = item.get('heart_score')
            heart_short = '—' if heart_val is None else _fmt(heart_val)
            self.telemetry_list.insert(
                tk.END,
                f"{item['t']:>6.1f}s | {item['class']} | ЧСС {heart_short} | Т {_fmt(item.get('telemetry_score'))} | ML {_fmt(item.get('ml_score'))} | И {_fmt(item.get('overall_score'))}",
            )

    def _draw_graph(self, cursor_t: Optional[float] = None):
        self.ax.clear()
        self.ax.grid(True, alpha=0.25)
        if not self.timeline:
            self.canvas.draw_idle()
            return
        xs = [float(r["t"]) for r in self.timeline]
        heart_series = [r["heart_score"] if r.get("heart_score") is not None else np.nan for r in self.timeline]
        self.ax.plot(xs, heart_series, label="ЧСС / HRV")
        self.ax.plot(xs, [r["telemetry_score"] for r in self.timeline], label="Телеметрия")
        if self.analysis_mode.get() == "ml":
            self.ax.plot(xs, [r["ml_score"] for r in self.timeline], label="ML")
        self.ax.plot(xs, [r["overall_score"] for r in self.timeline], label="Итог")
        if cursor_t is not None:
            self.ax.axvline(cursor_t, linestyle="--")
        self.ax.set_ylim(0, 1)
        self.ax.set_xlabel("Время, с")
        self.ax.legend(loc="upper left")
        self.fig.tight_layout()
        self.canvas.draw_idle()

    def calibrate_rule_based(self):
        try:
            data = self.app.api.post("/calibration/rule").json()
            self.app.rule_weights = data.get("weights", self.app.rule_weights)
            note = data.get("message", "Rule-based откалиброван")
            self.status.set(note)
            messagebox.showinfo("Калибровка", f"{note}\nНовые веса: {self.app.rule_weights}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось откалибровать rule-based: {e}")

    def train_ml(self):
        try:
            data = self.app.api.get("/research/train-model").json()
            msg = data.get("message", "Модель готова")
            messagebox.showinfo("ML", f"{msg}\nМодель: {data.get('model', '')}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обучить ML: {e}")

    def stop_session(self):
        if not self.running or not self.session_id:
            return
        self.running = False
        try:
            self.collector.stop()
        except Exception:
            pass
        video_path = None
        audio_path = None
        duration = max(1.0, max((self.timeline[-1]["t"] if self.timeline else 0.0), 1.0))
        out_dir = CACHE_DIR / f"session_{self.session_id}"
        out_dir.mkdir(parents=True, exist_ok=True)
        if self.recorder:
            video_path, audio_path = self.recorder.stop()
            duration = max(1.0, (self.recorder.end_ts - self.recorder.start_ts))
        else:
            video_path = create_blank_video(out_dir / "session.mp4", duration, fps=15, size=(1280, 720))
        if not audio_path:
            audio_path = synth_audio_from_timeline(self.timeline, out_dir / "session.wav", duration)
        last = self.timeline[-1] if self.timeline else {}
        summary = {
            "heart": last.get("heart_score", None),
            "telemetry": last.get("telemetry_score", 0.0),
            "ml": last.get("ml_score", 0.0),
            "overall": last.get("overall_score", 0.0),
            "class": last.get("class", "низкий"),
            "source": last.get("source", "Rule-based"),
            "metrics": last,
        }
        payload = {
            "timeline": self.timeline,
            "summary": summary,
            "heart_score": float(last.get("heart_score", 0.0)) if last.get("heart_score") is not None else 0.0,
            "telemetry_score": float(last.get("telemetry_score", 0.0)),
            "ml_score": float(last.get("ml_score", 0.0)),
            "overall_score": float(last.get("overall_score", 0.0)),
            "stress_class": str(last.get("class", "низкий")),
            "has_rr": bool(last.get("ibi_samples", 0)),
            "notes": self.notes.get().strip(),
            "ended_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "video_path": str(video_path or ""),
            "audio_path": str(audio_path or ""),
        }
        try:
            self.app.api.put(f"/sessions/{self.session_id}", json=payload)
        except Exception as e:
            messagebox.showwarning("Предупреждение", f"Сессия завершена локально, но не отправлена на сервер: {e}")
        if video_path and Path(video_path).exists():
            try:
                with open(video_path, "rb") as fh:
                    self.app.api.post(f"/sessions/{self.session_id}/video", files={"file": (Path(video_path).name, fh, "video/mp4")})
            except Exception:
                pass
        if audio_path and Path(audio_path).exists():
            try:
                with open(audio_path, "rb") as fh:
                    self.app.api.post(f"/sessions/{self.session_id}/audio", files={"file": (Path(audio_path).name, fh, "audio/wav")})
            except Exception:
                pass
        self.status.set(f"Сессия №{self.session_id} завершена")
        self.session_id = None
        self.recorder = None
        self.waiting = False
        self.running = False
        self.game_title.set("")  # Сбрасываем жесткую привязку к окну
        self._update_ui()
        self.app.refresh_views()


class PlaybackPanel(ttk.Frame):
    def __init__(self, master, app):
        super().__init__(master, padding=8)
        self.app = app
        self.sessions: List[Dict[str, Any]] = []
        self.current_session: Dict[str, Any] = {}
        self.video_path: Optional[Path] = None
        self.cap: Optional[cv2.VideoCapture] = None
        self.playing = False
        self.after_id = None
        self.duration = 0.0
        self.current_t = 0.0
        self.timeline: List[Dict[str, Any]] = []

        left = ttk.Frame(self)
        left.pack(side="left", fill="y", padx=(0, 10))
        right = ttk.Frame(self)
        right.pack(side="left", fill="both", expand=True)

        ttk.Label(left, text="Сессии", font=("Segoe UI", 11, "bold")).pack(anchor="w")
        self.session_list = tk.Listbox(left, width=40, height=25)
        self.session_list.pack(fill="both", expand=True)
        self.session_list.bind("<<ListboxSelect>>", lambda e: self.open_selected())
        ttk.Button(left, text="Обновить", command=self.refresh).pack(fill="x", pady=2)
        ttk.Button(left, text="Открыть", command=self.open_selected).pack(fill="x", pady=2)

        top = ttk.Frame(right)
        top.pack(fill="x")
        ttk.Button(top, text="▶", width=4, command=self.play).pack(side="left")
        ttk.Button(top, text="⏸", width=4, command=self.pause).pack(side="left", padx=(4, 10))
        self.slider = tk.DoubleVar(value=0.0)
        self.scale = ttk.Scale(top, from_=0.0, to=1.0, orient="horizontal", variable=self.slider, command=self._on_seek)
        self.scale.pack(side="left", fill="x", expand=True)
        self.time_label = ttk.Label(top, text="0.0 / 0.0")
        self.time_label.pack(side="left", padx=8)

        self.video_label = ttk.Label(right, text="Видео появится после выбора сессии", anchor="center")
        self.video_label.pack(fill="both", expand=False)

        self.telemetry_box = tk.Listbox(right, height=8)
        self.telemetry_box.pack(fill="both", expand=False, pady=(6, 6))

        self.fig = Figure(figsize=(6, 3.3), dpi=100)
        self.ax = self.fig.add_subplot(111)
        self.canvas = FigureCanvasTkAgg(self.fig, master=right)
        self.canvas.get_tk_widget().pack(fill="both", expand=True)

    def refresh(self):
        if not self.winfo_exists():
            return
        try:
            self.sessions = self.app.api.get("/sessions").json()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить сессии: {e}")
            return
        try:
            self.session_list.delete(0, tk.END)
            for s in self.sessions:
                self.session_list.insert(tk.END, f"№{s['id']} | {s['game_title']} | {s.get('player_username','')} | {s.get('stress_class','')}")
        except Exception:
            return

    def open_selected(self):
        idx = self.session_list.curselection()
        if not idx:
            return
        session = self.sessions[idx[0]]
        self.load_session(session)

    def load_session(self, session: Dict[str, Any]):
        self.current_session = session
        try:
            detail = self.app.api.get(f"/sessions/{session['id']}").json()
            self.timeline = detail.get("timeline", []) or []
            video_path = detail.get("video_path", "")
            self.video_path = None
            if video_path:
                remote = self.app.api.base_url + "/media/" + Path(video_path).name
                local = CACHE_DIR / f"session_{session['id']}.mp4"
                local.write_bytes(requests.get(remote, timeout=120).content)
                self.video_path = local
            self._prepare_video()
            self.render_at(0.0)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть сессию: {e}")

    def _prepare_video(self):
        if self.cap:
            self.cap.release()
        self.cap = None
        self.duration = 0.0
        if self.video_path and self.video_path.exists():
            self.cap = cv2.VideoCapture(str(self.video_path))
            fps = self.cap.get(cv2.CAP_PROP_FPS) or 5.0
            frames = self.cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0.0
            self.duration = float(frames / fps) if fps > 0 else 0.0
            self.scale.configure(to=max(self.duration, 1.0))
            self.time_label.configure(text=f"0.0 / {self.duration:.1f}")
        else:
            self.video_label.configure(image="", text="Видео не записывалось для этой сессии")
            self.video_label.image = None
            self.scale.configure(to=1.0)
            self.time_label.configure(text="0.0 / 0.0")

    def _on_seek(self, _value):
        try:
            self.render_at(float(self.slider.get()))
        except Exception:
            pass

    def play(self):
        if not self.cap:
            return
        self.playing = True
        self._tick()

    def pause(self):
        self.playing = False
        if self.after_id:
            self.after_cancel(self.after_id)
            self.after_id = None

    def _tick(self):
        if not self.playing or not self.cap:
            return
        fps = self.cap.get(cv2.CAP_PROP_FPS) or 30.0
        nxt = min(float(self.slider.get()) + 1.0 / max(fps, 1.0), self.duration)
        self.render_at(nxt)
        if nxt < self.duration:
            # Вычисляем задержку на основе FPS (для 30 FPS это ~33 мс)
            delay = int(1000.0 / max(fps, 1.0))
            self.after_id = self.after(delay, self._tick)

    def render_at(self, t: float):
        self.current_t = max(0.0, t)
        self.slider.set(self.current_t)
        self.time_label.configure(text=f"{self.current_t:.1f} / {self.duration:.1f}")
        self._show_frame(self.current_t)
        self._update_telemetry(self.current_t)
        self._draw_graph(self.current_t)

    def _show_frame(self, t: float):
        if not self.cap:
            return
        fps = self.cap.get(cv2.CAP_PROP_FPS) or 5.0
        self.cap.set(cv2.CAP_PROP_POS_FRAMES, int(t * fps))
        ok, frame = self.cap.read()
        if not ok:
            return
        img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
        img.thumbnail((980, 560))
        photo = ImageTk.PhotoImage(img)
        self.video_label.configure(image=photo, text="")
        self.video_label.image = photo

    def _update_telemetry(self, t: float):
        self.telemetry_box.delete(0, tk.END)
        for row in self.timeline:
            if float(row.get("t", 0.0)) <= t:
                self.telemetry_box.insert(
                    tk.END,
                    f"{row['t']:.1f}s | ЧСС {_fmt(row.get('heart_score'))} | Т {_fmt(row.get('telemetry_score'))} | ML {_fmt(row.get('ml_score'))} | И {_fmt(row.get('overall_score'))}",
                )
        self.telemetry_box.see(tk.END)

    def _draw_graph(self, cursor_t: Optional[float] = None):
        self.ax.clear()
        self.ax.grid(True, alpha=0.25)
        if not self.timeline:
            self.canvas.draw_idle()
            return
        xs = [float(r.get("t", 0.0)) for r in self.timeline]
        # Безопасное извлечение с заменой None на np.nan для корректной отрисовки matplotlib
        self.ax.plot(xs,
                     [float(r["heart_score"]) if r.get("heart_score") is not None else np.nan for r in self.timeline],
                     label="ЧСС / HRV")
        self.ax.plot(xs, [float(r["telemetry_score"]) if r.get("telemetry_score") is not None else 0.0 for r in
                          self.timeline], label="Телеметрия")
        self.ax.plot(xs,
                     [float(r["overall_score"]) if r.get("overall_score") is not None else 0.0 for r in self.timeline],
                     label="Итог")
        if cursor_t is not None:
            self.ax.axvline(cursor_t, linestyle="--")
        self.ax.set_ylim(0, 1)
        self.ax.set_xlabel("t, с")
        self.ax.legend(loc="upper left")
        self.fig.tight_layout()
        self.canvas.draw_idle()


class ResearchTab(ttk.Frame):
    def __init__(self, master, app):
        super().__init__(master, padding=8)
        self.app = app
        self.games: List[Dict[str, Any]] = []
        self.players: List[Dict[str, Any]] = []
        self.sessions: List[Dict[str, Any]] = []
        self.selected_game = tk.StringVar(value="Все игры")

        left = ttk.Frame(self)
        left.pack(side="left", fill="y", padx=(0, 8))
        right = ttk.Notebook(self)
        right.pack(side="left", fill="both", expand=True)

        self.games_box = ttk.Frame(right)
        self.players_box = ttk.Frame(right)
        self.sessions_box = ttk.Frame(right)
        self.compare_box = ttk.Frame(right)
        right.add(self.games_box, text="Игры")
        right.add(self.players_box, text="Игроки")
        right.add(self.sessions_box, text="Сессии")
        right.add(self.compare_box, text="Сравнение")

        ttk.Label(left, text="Действия", font=("Segoe UI", 11, "bold")).pack(anchor="w")
        ttk.Button(left, text="Обновить", command=self.refresh).pack(fill="x", pady=2)
        ttk.Button(left, text="Калибровать rule-based", command=self.calibrate_rule_based).pack(fill="x", pady=2)
        ttk.Button(left, text="Обучить ML", command=self.train_ml).pack(fill="x", pady=2)
        ttk.Button(left, text="Открыть выделенную сессию", command=self.open_selected_session).pack(fill="x", pady=2)
        ttk.Button(left, text="Выгрузить все таблицы в Word", command=self.export_all_tables_docx).pack(fill="x", pady=2)
        ttk.Button(left, text="Выгрузить сессии выбранной игры в Word", command=self.export_selected_game_tables_docx).pack(fill="x", pady=2)

        filter_frame = ttk.Frame(left)
        filter_frame.pack(fill="x", pady=(10, 2))
        ttk.Label(filter_frame, text="Фильтр оценки:").pack(side="left")
        self.session_filter = tk.StringVar(value="Все")
        filter_cb = ttk.Combobox(filter_frame, textvariable=self.session_filter, values=["Все", "Rule-based", "ML"],
                                 state="readonly", width=12)
        filter_cb.pack(side="right")
        self.session_filter.trace_add("write", lambda *args: self._fill_trees())

        game_filter_frame = ttk.Frame(left)
        game_filter_frame.pack(fill="x", pady=(8, 2))
        ttk.Label(game_filter_frame, text="Игра для таблиц:").pack(side="left")
        self.game_filter_cb = ttk.Combobox(game_filter_frame, textvariable=self.selected_game, values=["Все игры"], state="readonly", width=18)
        self.game_filter_cb.pack(side="right")
        self.selected_game.trace_add("write", lambda *args: self._fill_trees())

        self.game_list = tk.Listbox(left, width=38, height=10)
        self.game_list.pack(fill="both", expand=False, pady=(10, 0))
        self.game_list.bind("<<ListboxSelect>>", lambda e: self._on_game_selected())

        self.player_list = tk.Listbox(left, width=38, height=14)
        self.player_list.pack(fill="both", expand=True, pady=(8, 0))
        self.player_list.bind("<<ListboxSelect>>", lambda e: self.show_player_sessions())

        self._build_games_tab()
        self._build_players_tab()
        self._build_sessions_tab()
        self._build_compare_tab()

    def _tree(self, parent, columns: List[str], widths: Optional[List[int]] = None):
        tree = ttk.Treeview(parent, columns=columns, show="headings", height=12)
        for i, col in enumerate(columns):
            tree.heading(col, text=col)
            tree.column(col, width=(widths[i] if widths else 120), anchor="w")
        return tree

    def _current_game_title(self) -> str:
        value = (self.selected_game.get() or "").strip()
        return value if value and value != "Все игры" else ""

    def _session_rows_for_export(self) -> List[Dict[str, Any]]:
        title = self._current_game_title()
        if not title:
            return list(self.sessions)
        return [s for s in self.sessions if s.get("game_title", "") == title]

    def _on_game_selected(self):
        try:
            sel = self.game_list.curselection()
            if sel:
                title = self.game_list.get(sel[0])
                if title:
                    self.selected_game.set(title)
                    self.show_game_comparison()
                    self._fill_trees()
        except Exception:
            pass

    def _write_research_docx(self, out: str, sessions: List[Dict[str, Any]], title_suffix: str = ""):
        if Document is None:
            messagebox.showerror("Ошибка", "Для экспорта в Word требуется пакет python-docx")
            return
        if not self.games and not self.players and not self.sessions:
            self.refresh()
        doc = Document()
        heading = "Таблицы исследователя" if not title_suffix else f"Таблицы исследователя — {title_suffix}"
        doc.add_heading(heading, level=1)
        doc.add_paragraph(f"Сформировано: {time.strftime('%Y-%m-%d %H:%M:%S')}")

        # --- Таблица игр (строится по переданным сессиям) ---
        game_stats: Dict[str, Dict[str, Any]] = {}
        for s in sessions:
            title_raw = s.get("game_title", "")
            # Очищаем от переносов строк и лишних пробелов
            title = " ".join(title_raw.split())
            if not title:
                continue
            if title not in game_stats:
                game_stats[title] = {"sessions": 0, "players": set()}
            game_stats[title]["sessions"] += 1
            player = s.get("player_username") or s.get("player_id")
            if player:
                game_stats[title]["players"].add(str(player))

        game_rows = []
        for title, stats in game_stats.items():
            game_rows.append([title, stats["sessions"], len(stats["players"])])

        def add_table(title: str, headers: List[str], rows: List[List[Any]]):
            doc.add_heading(title, level=2)
            if not rows:
                doc.add_paragraph("Нет данных.")
                return
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, head in enumerate(headers):
                hdr[i].text = str(head)
            for row in rows:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)

        add_table("Игры", ["Игра", "Сессий", "Игроков"], game_rows)

        def add_table(title: str, headers: List[str], rows: List[List[Any]]):
            doc.add_heading(title, level=2)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, head in enumerate(headers):
                hdr[i].text = str(head)
            for row in rows:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)


        player_rows = []
        for p in self.players:
            prof = p.get("profile", {}) or {}
            player_rows.append([
                p.get("id", ""),
                p.get("username", ""),
                p.get("age", prof.get("age", 0)),
                p.get("sex", prof.get("sex", "")),
                p.get("disease_group", classify_disease_group(prof)),
                p.get("disease_code", disease_codes_for_profile(prof, 1)),
                p.get("group_index", ""),
            ])
        add_table("Игроки", ["ID", "Игрок", "Возраст", "Пол", "Группа", "Код", "Индекс"], player_rows)

        session_rows = []
        for s in sessions:
            session_rows.append([
                s.get("id", ""),
                s.get("game_title", ""),
                s.get("player_username", ""),
                s.get("age", ""),
                s.get("sex", ""),
                s.get("disease_group", ""),
                s.get("heart_score", ""),
                s.get("telemetry_score", ""),
                s.get("ml_score", ""),
                s.get("overall_score", ""),
                s.get("stress_class", ""),
            ])
        add_table("Сессии", ["ID", "Игра", "Игрок", "Возраст", "Пол", "Группа", "ЧСС", "Телеметрия", "ML", "Итог", "Класс"], session_rows)

        doc.save(out)
        messagebox.showinfo("Word", f"Таблицы сохранены в {out}")

    def export_all_tables_docx(self):
        default_name = f"research_tables_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        out = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=default_name, filetypes=[("Word document", "*.docx")])
        if not out:
            return
        self._write_research_docx(out, self._session_rows_for_export(), title_suffix=self._current_game_title() or "все игры")

    def export_selected_game_tables_docx(self):
        title = self._current_game_title()
        if not title:
            messagebox.showinfo("Word", "Сначала выберите игру для экспорта.")
            return
        default_name = f"research_tables_{slugify(title)}_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        out = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=default_name, filetypes=[("Word document", "*.docx")])
        if not out:
            return
        self._write_research_docx(out, [s for s in self.sessions if s.get("game_title", "") == title], title_suffix=title)

    def _build_games_tab(self):
        self.games_tree = self._tree(self.games_box, ["Игра", "Сессий", "Игроков"])
        self.games_tree.pack(fill="both", expand=True)

    def _build_players_tab(self):
        self.players_tree = self._tree(self.players_box, ["ID", "Игрок", "Возраст", "Пол", "Группа", "Код", "Сессий"], [60, 180, 70, 80, 120, 60, 70])
        self.players_tree.pack(fill="both", expand=True)
        self.players_tree.bind("<<TreeviewSelect>>", lambda e: self._select_player_from_tree())

    def _build_sessions_tab(self):
        self.sessions_tree = self._tree(self.sessions_box, ["ID", "Игра", "Игрок", "ЧСС", "Телеметрия", "ML", "Итог", "Класс"], [50, 160, 130, 70, 90, 70, 70, 80])
        self.sessions_tree.pack(fill="both", expand=True)
        self.sessions_tree.bind("<<TreeviewSelect>>", lambda e: self._select_session_from_tree())

    def _build_compare_tab(self):
        self.compare_text = tk.Text(self.compare_box, wrap="word")
        self.compare_text.pack(fill="both", expand=True)

    def refresh(self):
        if not self.winfo_exists():
            return
        try:
            self.games = self.app.api.get("/games").json()
            self.players = self.app.api.get("/players").json() if self.app.current_user.get("role") == ROLE_RESEARCHER else [self.app.api.get("/me").json()]
            self.sessions = self.app.api.get("/sessions").json()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обновить данные: {e}")
            return
        self._fill_lists()
        self._fill_trees()
        self.show_game_comparison()

    def _fill_lists(self):
        try:
            if self.game_list.winfo_exists():
                self.game_list.delete(0, tk.END)
                titles = ["Все игры"]
                for g in self.games:
                    title = g.get("title") if isinstance(g, dict) else str(g)
                    self.game_list.insert(tk.END, title)
                    titles.append(title)
                if hasattr(self, "game_filter_cb") and self.game_filter_cb.winfo_exists():
                    self.game_filter_cb["values"] = titles
                    if self.selected_game.get() not in titles:
                        self.selected_game.set("Все игры")
            if self.player_list.winfo_exists():
                self.player_list.delete(0, tk.END)
                for p in self.players:
                    prof = p.get("profile", {}) or {}
                    group = p.get("disease_group") or classify_disease_group(prof)
                    code = p.get("disease_code") or disease_codes_for_profile(prof or {"disease_group": group}, 1)
                    age = p.get("age", prof.get("age", 0))
                    sex = p.get("sex", prof.get("sex", ""))
                    self.player_list.insert(tk.END, f"#{p['id']} {p['username']} | {group} | {code} | {age} лет | {sex}")
        except tk.TclError:
            return
        except Exception:
            return

    def _fill_trees(self):
        try:
            for tr in (self.games_tree, self.players_tree, self.sessions_tree):
                if not tr.winfo_exists():
                    continue
                for item in tr.get_children():
                    tr.delete(item)
        except tk.TclError:
            return
        except Exception:
            return
        game_counts: Dict[str, int] = {}
        game_players: Dict[str, set] = {}
        for s in self.sessions:
            game_counts[s["game_title"]] = game_counts.get(s["game_title"], 0) + 1
            game_players.setdefault(s["game_title"], set()).add(s.get("player_username", ""))
        for g in self.games:
            title = g.get("title") if isinstance(g, dict) else str(g)
            self.games_tree.insert("", tk.END, values=(title, game_counts.get(title, 0), len(game_players.get(title, set()))))
        player_sessions = {p["id"]: 0 for p in self.players}
        for s in self.sessions:
            player_sessions[s.get("player_id")] = player_sessions.get(s.get("player_id"), 0) + 1
        for p in self.players:
            prof = p.get("profile", {}) or {}
            group = p.get("disease_group") or classify_disease_group(prof)
            age = p.get("age", prof.get("age", 0))
            sex = p.get("sex", prof.get("sex", ""))
            self.players_tree.insert(
                "",
                tk.END,
                iid=str(p["id"]),
                values=(p["id"], p["username"], age, sex, group, p.get("disease_code", disease_codes_for_profile(prof or {"disease_group": group}, 1)), player_sessions.get(p["id"], 0)),
            )
        mode_filter = self.session_filter.get()
        game_filter = self._current_game_title()
        for s in self.sessions:
            summary = safe_json_loads(s.get("summary_json") or "{}")
            source = summary.get("source", "Rule-based")

            # Пропускаем сессии, если они не подходят под выбранную игру/фильтр
            if mode_filter != "Все" and source != mode_filter:
                continue
            if game_filter and s.get("game_title", "") != game_filter:
                continue

            self.sessions_tree.insert(
                "",
                tk.END,
                iid=str(s["id"]),
                values=(
                    s["id"],
                    s.get("game_title", ""),
                    s.get("player_username", ""),
                    _fmt(s.get('heart_score')),
                    _fmt(s.get('telemetry_score')),
                    _fmt(s.get('ml_score')),
                    _fmt(s.get('overall_score')),
                    s.get("stress_class", ""),
                ),
            )

    def _select_player_from_tree(self):
        sel = self.players_tree.selection()
        if not sel:
            return
        player_id = int(sel[0])
        self.load_player_sessions(player_id)

    def _select_session_from_tree(self):
        sel = self.sessions_tree.selection()
        if not sel:
            return
        sid = int(sel[0])
        session = next((s for s in self.sessions if int(s.get("id", 0)) == sid), None)
        if session:
            self.app.open_session_in_player_view(session)

    def show_game_comparison(self):
        title = self._current_game_title()
        if not title:
            idx = self.game_list.curselection()
            if idx:
                title = self.game_list.get(idx[0])
                self.selected_game.set(title)
        if not title:
            self.compare_text.delete("1.0", tk.END)
            self.compare_text.insert(tk.END, "Выберите игру для сравнения.")
            return
        try:
            data = self.app.api.get(f"/research/comparison/{title}").json()
        except Exception as e:
            self.compare_text.delete("1.0", tk.END)
            self.compare_text.insert(tk.END, f"Не удалось получить сравнение: {e}")
            return
        groups = data.get("groups", {})
        lines = [f"Сравнение по игре: {title}\n"]
        healthy = groups.get("healthy", {}).get("overall_avg", 0.0) or 1.0
        for key, label in [("healthy", "Здоровые"), ("asthma", "Астма/лёгкие"), ("heart", "Сердечные")]:
            g = groups.get(key, {})
            lines.append(
                f"{label}: n={g.get('count', 0)}, итог {g.get('overall_avg', 0):.2f}, "
                f"ЧСС {g.get('heart_avg', 0):.2f}, телеметрия {g.get('telemetry_avg', 0):.2f}, ML {g.get('ml_avg', 0):.2f}, "
                f"к здоровым {g.get('vs_healthy_percent', 0):+.1f}%"
            )
        self.compare_text.delete("1.0", tk.END)
        self.compare_text.insert(tk.END, "\n".join(lines))

    def show_player_sessions(self):
        idx = self.player_list.curselection()
        if not idx:
            return
        player = self.players[idx[0]]
        self.load_player_sessions(int(player["id"]))

    def load_player_sessions(self, player_id: int):
        try:
            sessions = self.app.api.get(f"/research/players/{player_id}/sessions").json()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить сессии игрока: {e}")
            return
        self.sessions_tree.delete(*self.sessions_tree.get_children())
        game_filter = self._current_game_title()
        for s in sessions:
            if game_filter and s.get("game_title", "") != game_filter:
                continue
            self.sessions_tree.insert(
                "",
                tk.END,
                iid=str(s["id"]),
                values=(
                    s["id"],
                    s.get("game_title", ""),
                    s.get("player_username", ""),
                    _fmt(s.get('heart_score')),
                    _fmt(s.get('telemetry_score')),
                    _fmt(s.get('ml_score')),
                    _fmt(s.get('overall_score')),
                    s.get("stress_class", ""),
                ),
            )
        self.app.last_research_player_id = player_id
        self.compare_text.delete("1.0", tk.END)
        self.compare_text.insert(tk.END, f"Выбран игрок #{player_id}. Сессии загружены.\n")

    def open_selected_session(self):
        sel = self.sessions_tree.selection()
        if not sel:
            return
        sid = int(sel[0])
        session = next((s for s in self.sessions if int(s.get("id", 0)) == sid), None)
        if session:
            self.app.open_session_in_player_view(session)

    def calibrate_rule_based(self):
        try:
            data = self.app.api.post("/calibration/rule").json()
            self.app.rule_weights = data.get("weights", self.app.rule_weights)
            messagebox.showinfo("Rule-based", f"Новые веса: {self.app.rule_weights}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось калибровать rule-based: {e}")

    def train_ml(self):
        try:
            data = self.app.api.get("/research/train-model").json()
            messagebox.showinfo("ML", f"Модель обучена: {data.get('model', '')}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обучить ML: {e}")


class StressDesktopApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_NAME)
        self.geometry("1500x960")
        self.minsize(1280, 820)
        self.api = ApiClient(DEFAULT_API_URL)
        self.current_user: Dict[str, Any] = {}
        self.rule_weights = load_rule_calibration().get("weights", {"heart": 0.40, "telemetry": 0.35, "ml": 0.25})
        self.token_path = CACHE_DIR / "token.json"
        self.last_research_player_id: Optional[int] = None

        self.style = ttk.Style(self)
        try:
            self.style.theme_use("clam")
        except Exception:
            pass
        self.style.configure("TButton", padding=6)
        self.style.configure("TLabel", padding=2)
        self.style.configure("Treeview", rowheight=24)

        self.container = ttk.Frame(self, padding=10)
        self.container.pack(fill="both", expand=True)
        self.login_frame = LoginFrame(self.container, self)
        self.login_frame.pack(fill="both", expand=True)

        self.main_notebook: Optional[ttk.Notebook] = None
        self.live_tab: Optional[LivePlayerTab] = None
        self.playback_tab: Optional[PlaybackPanel] = None
        self.research_tab: Optional[ResearchTab] = None

        self._load_saved_login()

    def _load_saved_login(self):
        if not self.token_path.exists():
            return
        try:
            data = json.loads(self.token_path.read_text(encoding="utf-8"))
            token = data.get("token", "")
            user = data.get("user", {})
            if token:
                self.api.set_auth(token, user)
                self.current_user = user
                self.show_main()
        except Exception:
            pass

    def save_login(self, token: str, user: Dict[str, Any], remember: bool):
        self.api.set_auth(token, user)
        self.current_user = user
        if remember:
            self.token_path.write_text(json.dumps({"token": token, "user": user}, ensure_ascii=False, indent=2), encoding="utf-8")
        self.show_main()

    def refresh_rule_weights(self):
        try:
            data = self.api.get("/calibration/rule").json()
            weights = data.get("weights")
            if isinstance(weights, dict):
                self.rule_weights = {
                    "heart": float(weights.get("heart", self.rule_weights.get("heart", 0.40))),
                    "telemetry": float(weights.get("telemetry", self.rule_weights.get("telemetry", 0.35))),
                    "ml": float(weights.get("ml", self.rule_weights.get("ml", 0.25))),
                }
                return
        except Exception:
            pass
        try:
            self.rule_weights = load_rule_calibration().get("weights", self.rule_weights)
        except Exception:
            pass

    def do_login(self, username: str, password: str):
        try:
            data = self.api.post("/auth/login", json={"username": username, "password": password}).json()
            self.save_login(data["token"], data["user"], self.login_frame.remember.get())
        except Exception as e:
            messagebox.showerror("Ошибка входа", str(e))

    def do_register(self, username: str, password: str, role: str):
        try:
            data = self.api.post("/auth/register", json={"username": username, "password": password, "role": role}).json()
            self.save_login(data["token"], data["user"], self.login_frame.remember.get())
        except Exception as e:
            messagebox.showerror("Ошибка регистрации", str(e))

    def show_main(self):
        self.refresh_rule_weights()
        for child in self.container.winfo_children():
            child.destroy()
        head = ttk.Frame(self.container)
        head.pack(fill="x", pady=(0, 8))
        ttk.Label(head, text=f"Пользователь: {self.current_user.get('username', '')}", font=("Segoe UI", 12, "bold")).pack(side="left")
        ttk.Label(head, text=f"Роль: {self.current_user.get('role_label', '')}", font=("Segoe UI", 12)).pack(side="left", padx=18)
        ttk.Button(head, text="Выйти", command=self.logout).pack(side="right")

        self.main_notebook = ttk.Notebook(self.container)
        self.main_notebook.pack(fill="both", expand=True)

        if self.current_user.get("role") == ROLE_PLAYER:
            self.live_tab = LivePlayerTab(self.main_notebook, self)
            self.playback_tab = PlaybackPanel(self.main_notebook, self)
            self.main_notebook.add(self.live_tab, text="В реальном времени")
            self.main_notebook.add(self.playback_tab, text="Сессии и результаты")
            self.playback_tab.refresh()
        else:
            self.research_tab = ResearchTab(self.main_notebook, self)
            self.playback_tab = PlaybackPanel(self.main_notebook, self)
            self.main_notebook.add(self.research_tab, text="Игры, игроки и сравнение")
            self.main_notebook.add(self.playback_tab, text="Просмотр сессий")
            self.research_tab.refresh()
            self.playback_tab.refresh()

    def logout(self):
        self.api.set_auth("", {})
        self.current_user = {}
        try:
            if self.token_path.exists():
                self.token_path.unlink()
        except Exception:
            pass
        for child in self.container.winfo_children():
            child.destroy()
        self.login_frame = LoginFrame(self.container, self)
        self.login_frame.pack(fill="both", expand=True)

    def refresh_views(self):
        for tab in (self.research_tab, self.playback_tab):
            try:
                if tab is not None and tab.winfo_exists():
                    tab.after(0, tab.refresh)
            except Exception:
                pass

    def open_session_in_player_view(self, session: Dict[str, Any]):
        if self.playback_tab is None:
            return
        if self.main_notebook:
            try:
                self.main_notebook.select(self.playback_tab)
            except Exception:
                pass
        self.playback_tab.load_session(session)


def main():
    app = StressDesktopApp()
    app.mainloop()


if __name__ == "__main__":
    main()
